#!/usr/bin/env python3
"""
generate_report.py

Generates a .docx plagiarism report from a Markdown template and a JSON data file.

Usage:
    python generate_report.py <template.md> <data.json>

The .docx is saved alongside the .json config file with the same base name.
Template variables use ${key} notation; keys must match the top-level keys in
the JSON file. A bare '$' in the markdown must be escaped as '$$'.

Supported Markdown:
    # / ## / ###     Headings
    **text**         Bold
    *text*           Italic
    <u>text</u>      Underline
    `code`           Inline code (gray monospace)
    ```...```        Code block (gray monospace, preserves line breaks)
    - or * item      Bullet list
    1. item          Ordered list
"""

import sys
import json
import re
from string import Template
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Inline formatting ──────────────────────────────────────────────────────────

# Order matters: ** must be matched before * to avoid consuming the first star.
INLINE_RE = re.compile(
    r'\*\*(.*?)\*\*'    # bold
    r'|\*(.*?)\*'        # italic
    r'|<u>(.*?)</u>'     # underline
    r'|`(.*?)`'          # inline code
    r'|([^*<`]+)'        # plain text chunk
    r'|([\s\S])'         # fallback: any single character not matched above
)


def _parse_inline(text):
    """
    Returns a list of (content, bold, italic, underline, is_code) tuples
    representing the inline-formatted segments of a line of text.
    """
    runs = []
    for m in INLINE_RE.finditer(text):
        bold, italic, under, code, plain, fallback = m.groups()
        if bold     is not None: runs.append((bold,     True,  False, False, False))
        elif italic  is not None: runs.append((italic,  False, True,  False, False))
        elif under   is not None: runs.append((under,   False, False, True,  False))
        elif code    is not None: runs.append((code,    False, False, False, True ))
        elif plain   is not None: runs.append((plain,   False, False, False, False))
        elif fallback is not None: runs.append((fallback, False, False, False, False))
    return runs


def _add_inline(paragraph, text):
    """Parse inline markdown and append formatted runs to an existing paragraph."""
    for content, bold, italic, underline, is_code in _parse_inline(text):
        run = paragraph.add_run(content)
        if is_code:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            _shade_run(run, 'EBEBEB')
        else:
            run.bold      = bold
            run.italic    = italic
            run.underline = underline


# ── XML shading helpers ────────────────────────────────────────────────────────

def _shade_paragraph(paragraph, fill_hex):
    """Apply a background fill color to an entire paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    pPr.append(shd)


def _shade_run(run, fill_hex):
    """Apply a background fill color to a single run (character-level shading)."""
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    rPr.append(shd)


# ── Block parser ───────────────────────────────────────────────────────────────

def _is_bullet(line):
    return bool(re.match(r'^[-*] ', line))

def _is_ordered(line):
    return bool(re.match(r'^\d+\. ', line))

def _is_block_boundary(line):
    return (
        line.strip() == ''
        or line.startswith('#')
        or line.strip().startswith('```')
        or _is_bullet(line)
        or _is_ordered(line)
    )


def parse_blocks(md_text):
    """
    Converts a markdown string into a list of block dicts. Each dict has a
    'type' key and type-specific content:

        {'type': 'heading',      'level': int, 'content': str}
        {'type': 'paragraph',    'content': str}
        {'type': 'bullet_list',  'items': [str, ...]}
        {'type': 'ordered_list', 'items': [str, ...]}
        {'type': 'code',         'content': str}
    """
    blocks = []
    lines  = md_text.split('\n')
    i      = 0

    while i < len(lines):
        line = lines[i]

        # ── Code block (may contain blank lines, handled before anything else)
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            blocks.append({'type': 'code', 'content': '\n'.join(code_lines)})

        # ── Headings
        elif line.startswith('### '):
            blocks.append({'type': 'heading', 'level': 3, 'content': line[4:].strip()})
            i += 1
        elif line.startswith('## '):
            blocks.append({'type': 'heading', 'level': 2, 'content': line[3:].strip()})
            i += 1
        elif line.startswith('# '):
            blocks.append({'type': 'heading', 'level': 1, 'content': line[2:].strip()})
            i += 1

        # ── Bullet list (collect consecutive bullet lines)
        elif _is_bullet(line):
            items = []
            while i < len(lines) and _is_bullet(lines[i]):
                items.append(lines[i][2:])
                i += 1
            blocks.append({'type': 'bullet_list', 'items': items})

        # ── Ordered list (collect consecutive numbered lines)
        elif _is_ordered(line):
            items = []
            while i < len(lines) and _is_ordered(lines[i]):
                items.append(re.sub(r'^\d+\. ', '', lines[i]))
                i += 1
            blocks.append({'type': 'ordered_list', 'items': items})

        # ── Blank line (separator, no block produced)
        elif line.strip() == '':
            i += 1

        # ── Paragraph (collect until the next blank line or block-level element)
        else:
            para_lines = []
            while i < len(lines) and not _is_block_boundary(lines[i]):
                para_lines.append(lines[i].strip())
                i += 1
            if para_lines:
                blocks.append({'type': 'paragraph', 'content': ' '.join(para_lines)})

    return blocks


# ── Document builder ───────────────────────────────────────────────────────────

def build_docx(blocks):
    doc = Document()

    # Change default font-sizes for normal, h1, h2, and h3.
    doc.styles['Normal'].font.size = Pt(12)
    doc.styles['Heading 1'].font.size = Pt(24)
    doc.styles['Heading 2'].font.size = Pt(18)
    doc.styles['Heading 3'].font.size = Pt(14)

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    for block in blocks:
        btype = block['type']

        if btype == 'heading':
            # H1: dark maroon centered; H2: medium maroon; H3: light maroon
            HEADING_COLORS = {
                1: RGBColor(0x50, 0x00, 0x00),
                2: RGBColor(0x7B, 0x2D, 0x2D),
                3: RGBColor(0xA6, 0x50, 0x50),
            }
            p = doc.add_heading('', level=block['level'])
            _add_inline(p, block['content'])
            color = HEADING_COLORS[block['level']]
            for run in p.runs:
                run.font.color.rgb = color

        elif btype == 'paragraph':
            p = doc.add_paragraph()
            _add_inline(p, block['content'])

        elif btype == 'bullet_list':
            for item in block['items']:
                p = doc.add_paragraph(style='List Bullet')
                _add_inline(p, item)

        elif btype == 'ordered_list':
            for item in block['items']:
                p = doc.add_paragraph(style='List Number')
                _add_inline(p, item)

        elif btype == 'code':
            p = doc.add_paragraph()
            _shade_paragraph(p, 'F0F0F0')
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            # Preserve newlines inside the block using line breaks within one paragraph.
            code_lines = block['content'].split('\n')
            for j, code_line in enumerate(code_lines):
                run = p.add_run(code_line)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                if j < len(code_lines) - 1:
                    run.add_break()

    return doc


# ── Entry point ────────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) != 3:
        print("Usage: python generate_report.py <template.md> <data.json>")
        sys.exit(1)

    md_path   = Path(sys.argv[1])
    json_path = Path(sys.argv[2])

    if not md_path.exists():
        print(f"Error: Markdown template not found: {md_path}")
        sys.exit(1)
    if not json_path.exists():
        print(f"Error: JSON data file not found: {json_path}")
        sys.exit(1)
    if md_path.suffix.lower() != '.md':
        print(f"Error: First argument must be a .md file, got: {md_path.name}")
        sys.exit(1)
    if json_path.suffix.lower() != '.json':
        print(f"Error: Second argument must be a .json file, got: {json_path.name}")
        sys.exit(1)

    # Load and validate JSON
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except json.JSONDecodeError as e:
        print(f"Error: Could not parse JSON file: {e}")
        sys.exit(1)

    if not isinstance(data, dict):
        print("Error: JSON file must be a top-level object (key-value dictionary).")
        sys.exit(1)
    nested = [k for k, v in data.items() if isinstance(v, (dict, list))]
    if nested:
        print(f"Error: JSON file must be a flat dictionary. "
              f"Nested values found at key(s): {', '.join(nested)}")
        sys.exit(1)

    # Convert all values to strings so Template.substitute always has str to work with
    data = {k: str(v) for k, v in data.items()}

    # Load markdown and substitute template variables
    raw_md = md_path.read_text(encoding='utf-8')

    try:
        md_text = Template(raw_md).substitute(data)
    except ValueError as e:
        print(f"Error: Invalid placeholder in template -- {e}")
        print("Hint: A lone '$' in your markdown must be written as '$$'.")
        sys.exit(1)
    except KeyError as e:
        print(f"Error: Template uses key {e} which is not present in the JSON file.")
        sys.exit(1)

    # Parse markdown and build the document
    blocks = parse_blocks(md_text)
    doc    = build_docx(blocks)

    output_path = json_path.with_suffix('.docx')
    doc.save(str(output_path))
    print(f"Report saved to: {output_path}")


if __name__ == '__main__':
    main()
